import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Optional, Union

from docling_core.types.doc import (
    DocItemLabel,
    DoclingDocument,
    DocumentOrigin,
    GroupLabel,
    ImageRef,
    NodeItem,
    TableCell,
    TableData,
)
from docling_core.types.doc.document import Formatting
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tc
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.table import Table, _Cell
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree
from lxml.etree import XPath
from PIL import Image, UnidentifiedImageError
from pydantic import AnyUrl
from typing_extensions import override

from docling.backend.abstract_backend import DeclarativeDocumentBackend
from docling.backend.docx.latex.omml import oMath2Latex
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument

_log = logging.getLogger(__name__)


class MsWordDocumentBackend(DeclarativeDocumentBackend):
    @override
    def __init__(
        self, in_doc: "InputDocument", path_or_stream: Union[BytesIO, Path]
    ) -> None:
        super().__init__(in_doc, path_or_stream)
        self.XML_KEY = (
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        )
        self.xml_namespaces = {
            "w": "http://schemas.microsoft.com/office/word/2003/wordml"
        }
        # self.initialise(path_or_stream)
        # Word file:
        self.path_or_stream: Union[BytesIO, Path] = path_or_stream
        self.valid: bool = False
        # Initialise the parents for the hierarchy
        self.max_levels: int = 10
        self.level_at_new_list: Optional[int] = None
        self.parents: dict[int, Optional[NodeItem]] = {}
        self.numbered_headers: dict[int, int] = {}
        self.equation_bookends: str = "<eq>{EQ}</eq>"
        for i in range(-1, self.max_levels):
            self.parents[i] = None

        self.level = 0
        self.listIter = 0

        self.history: dict[str, Any] = {
            "names": [None],
            "levels": [None],
            "numids": [None],
            "indents": [None],
        }

        self.docx_obj = None
        try:
            if isinstance(self.path_or_stream, BytesIO):
                self.docx_obj = Document(self.path_or_stream)
            elif isinstance(self.path_or_stream, Path):
                self.docx_obj = Document(str(self.path_or_stream))

            self.valid = True
        except Exception as e:
            raise RuntimeError(
                f"MsPowerpointDocumentBackend could not load document with hash {self.document_hash}"
            ) from e

    @override
    def is_valid(self) -> bool:
        return self.valid

    @classmethod
    @override
    def supports_pagination(cls) -> bool:
        return False

    @override
    def unload(self):
        if isinstance(self.path_or_stream, BytesIO):
            self.path_or_stream.close()

        self.path_or_stream = None

    @classmethod
    @override
    def supported_formats(cls) -> set[InputFormat]:
        return {InputFormat.DOCX}

    @override
    def convert(self) -> DoclingDocument:
        """Parses the DOCX into a structured document model.

        Returns:
            The parsed document.
        """

        origin = DocumentOrigin(
            filename=self.file.name or "file",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            binary_hash=self.document_hash,
        )

        doc = DoclingDocument(name=self.file.stem or "file", origin=origin)
        if self.is_valid():
            assert self.docx_obj is not None
            doc = self._walk_linear(self.docx_obj.element.body, self.docx_obj, doc)
            return doc
        else:
            raise RuntimeError(
                f"Cannot convert doc with {self.document_hash} because the backend failed to init."
            )

    def _update_history(
        self,
        name: str,
        level: Optional[int],
        numid: Optional[int],
        ilevel: Optional[int],
    ):
        self.history["names"].append(name)
        self.history["levels"].append(level)

        self.history["numids"].append(numid)
        self.history["indents"].append(ilevel)

    def _prev_name(self) -> Optional[str]:
        return self.history["names"][-1]

    def _prev_level(self) -> Optional[int]:
        return self.history["levels"][-1]

    def _prev_numid(self) -> Optional[int]:
        return self.history["numids"][-1]

    def _prev_indent(self) -> Optional[int]:
        return self.history["indents"][-1]

    def _get_level(self) -> int:
        """Return the first None index."""
        for k, v in self.parents.items():
            if k >= 0 and v == None:
                return k
        return 0

    def _walk_linear(
        self,
        body: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> DoclingDocument:
        for element in body:
            tag_name = etree.QName(element).localname
            # Check for Inline Images (blip elements)
            namespaces = {
                "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            }
            xpath_expr = XPath(".//a:blip", namespaces=namespaces)
            drawing_blip = xpath_expr(element)

            # Check for Tables
            if element.tag.endswith("tbl"):
                try:
                    self._handle_tables(element, docx_obj, doc)
                except Exception:
                    _log.debug("could not parse a table, broken docx table")

            elif drawing_blip:
                self._handle_pictures(docx_obj, drawing_blip, doc)
            # Check for the sdt containers, like table of contents
            elif tag_name in ["sdt"]:
                sdt_content = element.find(".//w:sdtContent", namespaces=namespaces)
                if sdt_content is not None:
                    # Iterate paragraphs, runs, or text inside <w:sdtContent>.
                    paragraphs = sdt_content.findall(".//w:p", namespaces=namespaces)
                    for p in paragraphs:
                        self._handle_text_elements(p, docx_obj, doc)
            # Check for Text
            elif tag_name in ["p"]:
                # "tcPr", "sectPr"
                self._handle_text_elements(element, docx_obj, doc)
            else:
                _log.debug(f"Ignoring element in DOCX with tag: {tag_name}")
        return doc

    def _str_to_int(
        self, s: Optional[str], default: Optional[int] = 0
    ) -> Optional[int]:
        if s is None:
            return None
        try:
            return int(s)
        except ValueError:
            return default

    def _split_text_and_number(self, input_string: str) -> list[str]:
        match = re.match(r"(\D+)(\d+)$|^(\d+)(\D+)", input_string)
        if match:
            parts = list(filter(None, match.groups()))
            return parts
        else:
            return [input_string]

    def _get_numId_and_ilvl(
        self, paragraph: Paragraph
    ) -> tuple[Optional[int], Optional[int]]:
        # Access the XML element of the paragraph
        numPr = paragraph._element.find(
            ".//w:numPr", namespaces=paragraph._element.nsmap
        )

        if numPr is not None:
            # Get the numId element and extract the value
            numId_elem = numPr.find("w:numId", namespaces=paragraph._element.nsmap)
            ilvl_elem = numPr.find("w:ilvl", namespaces=paragraph._element.nsmap)
            numId = numId_elem.get(self.XML_KEY) if numId_elem is not None else None
            ilvl = ilvl_elem.get(self.XML_KEY) if ilvl_elem is not None else None

            return self._str_to_int(numId, None), self._str_to_int(ilvl, None)

        return None, None  # If the paragraph is not part of a list

    def _get_heading_and_level(self, style_label: str) -> tuple[str, Optional[int]]:
        parts = self._split_text_and_number(style_label)

        if len(parts) == 2:
            parts.sort()
            label_str: str = ""
            label_level: Optional[int] = 0
            if parts[0].strip().lower() == "heading":
                label_str = "Heading"
                label_level = self._str_to_int(parts[1], None)
            if parts[1].strip().lower() == "heading":
                label_str = "Heading"
                label_level = self._str_to_int(parts[0], None)
            return label_str, label_level

        return style_label, None

    def _get_label_and_level(self, paragraph: Paragraph) -> tuple[str, Optional[int]]:
        if paragraph.style is None:
            return "Normal", None

        label = paragraph.style.style_id
        name = paragraph.style.name
        base_style_label = None
        base_style_name = None
        if base_style := getattr(paragraph.style, "base_style", None):
            base_style_label = base_style.style_id
            base_style_name = base_style.name

        if label is None:
            return "Normal", None

        if ":" in label:
            parts = label.split(":")
            if len(parts) == 2:
                return parts[0], self._str_to_int(parts[1], None)

        if "heading" in label.lower():
            return self._get_heading_and_level(label)
        if "heading" in name.lower():
            return self._get_heading_and_level(name)
        if base_style_label and "heading" in base_style_label.lower():
            return self._get_heading_and_level(base_style_label)
        if base_style_name and "heading" in base_style_name.lower():
            return self._get_heading_and_level(base_style_name)

        return label, None

    @classmethod
    def _get_format_from_run(cls, run: Run) -> Optional[Formatting]:
        has_any_formatting = run.bold or run.italic or run.underline
        return (
            Formatting(
                bold=run.bold or False,
                italic=run.italic or False,
                underline=run.underline or False,
            )
            if has_any_formatting
            else None
        )

    def _get_paragraph_elements(self, paragraph: Paragraph):
        """
        Extract paragraph elements along with their formatting and hyperlink
        """

        # for now retain empty paragraphs for backwards compatibility:
        if paragraph.text.strip() == "":
            return [("", None, None)]

        paragraph_elements: list[
            tuple[str, Optional[Formatting], Optional[Union[AnyUrl, Path]]]
        ] = []
        group_text = ""
        previous_format = None

        # Iterate over the runs of the paragraph and group them by format
        for c in paragraph.iter_inner_content():
            if isinstance(c, Hyperlink):
                text = c.text
                hyperlink = Path(c.address)
                format = self._get_format_from_run(c.runs[0])
            elif isinstance(c, Run):
                text = c.text
                hyperlink = None
                format = self._get_format_from_run(c)
            else:
                continue

            if (len(text.strip()) and format != previous_format) or (
                hyperlink is not None
            ):
                # If the style changes for a non empty text, add the previous group
                if len(group_text.strip()) > 0:
                    paragraph_elements.append(
                        (group_text.strip(), previous_format, None)
                    )
                group_text = ""

                # If there is a hyperlink, add it immediately
                if hyperlink is not None:
                    paragraph_elements.append((text.strip(), format, hyperlink))
                    text = ""
                else:
                    previous_format = format

            group_text += text

        # Format the last group
        if len(group_text.strip()) > 0:
            paragraph_elements.append((group_text.strip(), format, None))

        return paragraph_elements

    def _handle_equations_in_text(self, element, text):
        only_texts = []
        only_equations = []
        texts_and_equations = []
        for subt in element.iter():
            tag_name = etree.QName(subt).localname
            if tag_name == "t" and "math" not in subt.tag:
                if isinstance(subt.text, str):
                    only_texts.append(subt.text)
                    texts_and_equations.append(subt.text)
            elif "oMath" in subt.tag and "oMathPara" not in subt.tag:
                latex_equation = str(oMath2Latex(subt)).strip()
                if len(latex_equation) > 0:
                    only_equations.append(
                        self.equation_bookends.format(EQ=latex_equation)
                    )
                    texts_and_equations.append(
                        self.equation_bookends.format(EQ=latex_equation)
                    )

        if len(only_equations) < 1:
            return text, []

        if (
            re.sub(r"\s+", "", "".join(only_texts)).strip()
            != re.sub(r"\s+", "", text).strip()
        ):
            # If we are not able to reconstruct the initial raw text
            # do not try to parse equations and return the original
            return text, []

        # Insert equations into original text
        # This is done to preserve white space structure
        output_text = text[:]
        init_i = 0
        for i_substr, substr in enumerate(texts_and_equations):
            if len(substr) == 0:
                continue

            if substr in output_text[init_i:]:
                init_i += output_text[init_i:].find(substr) + len(substr)
            else:
                if i_substr > 0:
                    output_text = output_text[:init_i] + substr + output_text[init_i:]
                    init_i += len(substr)
                else:
                    output_text = substr + output_text

        return output_text, only_equations

    def _create_or_reuse_parent(
        self,
        *,
        doc: DoclingDocument,
        prev_parent: Optional[NodeItem],
        paragraph_elements: list,
    ) -> Optional[NodeItem]:
        return (
            doc.add_group(label=GroupLabel.INLINE, parent=prev_parent)
            if len(paragraph_elements) > 1
            else prev_parent
        )

    def _handle_text_elements(
        self,
        element: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> None:
        paragraph = Paragraph(element, docx_obj)

        raw_text = paragraph.text
        text, equations = self._handle_equations_in_text(element=element, text=raw_text)

        if text is None:
            return
        paragraph_elements = self._get_paragraph_elements(paragraph)
        text = text.strip()

        # Common styles for bullet and numbered lists.
        # "List Bullet", "List Number", "List Paragraph"
        # Identify wether list is a numbered list or not
        # is_numbered = "List Bullet" not in paragraph.style.name
        is_numbered = False
        p_style_id, p_level = self._get_label_and_level(paragraph)
        numid, ilevel = self._get_numId_and_ilvl(paragraph)

        if numid == 0:
            numid = None

        # Handle lists
        if (
            numid is not None
            and ilevel is not None
            and p_style_id not in ["Title", "Heading"]
        ):
            self._add_list_item(
                doc=doc,
                numid=numid,
                ilevel=ilevel,
                elements=paragraph_elements,
                is_numbered=is_numbered,
            )
            self._update_history(p_style_id, p_level, numid, ilevel)
            return
        elif (
            numid is None
            and self._prev_numid() is not None
            and p_style_id not in ["Title", "Heading"]
        ):  # Close list
            if self.level_at_new_list:
                for key in range(len(self.parents)):
                    if key >= self.level_at_new_list:
                        self.parents[key] = None
                self.level = self.level_at_new_list - 1
                self.level_at_new_list = None
            else:
                for key in range(len(self.parents)):
                    self.parents[key] = None
                self.level = 0

        if p_style_id in ["Title"]:
            for key in range(len(self.parents)):
                self.parents[key] = None
            self.parents[0] = doc.add_text(
                parent=None, label=DocItemLabel.TITLE, text=text
            )
        elif "Heading" in p_style_id:
            style_element = getattr(paragraph.style, "element", None)
            if style_element is not None:
                is_numbered_style = (
                    "<w:numPr>" in style_element.xml or "<w:numPr>" in element.xml
                )
            else:
                is_numbered_style = False
            self._add_header(doc, p_level, text, is_numbered_style)

        elif len(equations) > 0:
            if (raw_text is None or len(raw_text.strip()) == 0) and len(text) > 0:
                # Standalone equation
                level = self._get_level()
                doc.add_text(
                    label=DocItemLabel.FORMULA,
                    parent=self.parents[level - 1],
                    text=text.replace("<eq>", "").replace("</eq>", ""),
                )
            else:
                # Inline equation
                level = self._get_level()
                inline_equation = doc.add_group(
                    label=GroupLabel.INLINE, parent=self.parents[level - 1]
                )
                text_tmp = text
                for eq in equations:
                    if len(text_tmp) == 0:
                        break

                    split_text_tmp = text_tmp.split(eq.strip(), maxsplit=1)

                    pre_eq_text = split_text_tmp[0]
                    text_tmp = "" if len(split_text_tmp) == 1 else split_text_tmp[1]

                    if len(pre_eq_text) > 0:
                        doc.add_text(
                            label=DocItemLabel.PARAGRAPH,
                            parent=inline_equation,
                            text=pre_eq_text,
                        )
                    doc.add_text(
                        label=DocItemLabel.FORMULA,
                        parent=inline_equation,
                        text=eq.replace("<eq>", "").replace("</eq>", ""),
                    )

                if len(text_tmp) > 0:
                    doc.add_text(
                        label=DocItemLabel.PARAGRAPH,
                        parent=inline_equation,
                        text=text_tmp.strip(),
                    )

        elif p_style_id in [
            "Paragraph",
            "Normal",
            "Subtitle",
            "Author",
            "DefaultText",
            "ListParagraph",
            "ListBullet",
            "Quote",
        ]:
            level = self._get_level()
            parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents.get(level - 1),
                paragraph_elements=paragraph_elements,
            )
            for text, format, hyperlink in paragraph_elements:
                doc.add_text(
                    label=DocItemLabel.PARAGRAPH,
                    parent=parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )

        else:
            # Text style names can, and will have, not only default values but user values too
            # hence we treat all other labels as pure text
            level = self._get_level()
            parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents.get(level - 1),
                paragraph_elements=paragraph_elements,
            )
            for text, format, hyperlink in paragraph_elements:
                doc.add_text(
                    label=DocItemLabel.PARAGRAPH,
                    parent=parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )

        self._update_history(p_style_id, p_level, numid, ilevel)
        return

    def _add_header(
        self,
        doc: DoclingDocument,
        curr_level: Optional[int],
        text: str,
        is_numbered_style: bool = False,
    ) -> None:
        level = self._get_level()
        if isinstance(curr_level, int):
            if curr_level > level:
                # add invisible group
                for i in range(level, curr_level):
                    self.parents[i] = doc.add_group(
                        parent=self.parents[i - 1],
                        label=GroupLabel.SECTION,
                        name=f"header-{i}",
                    )
            elif curr_level < level:
                # remove the tail
                for key in range(len(self.parents)):
                    if key >= curr_level:
                        self.parents[key] = None

            current_level = curr_level
            parent_level = curr_level - 1
            add_level = curr_level
        else:
            current_level = self.level
            parent_level = self.level - 1
            add_level = 1

        if is_numbered_style:
            if add_level in self.numbered_headers:
                self.numbered_headers[add_level] += 1
            else:
                self.numbered_headers[add_level] = 1
            text = f"{self.numbered_headers[add_level]} {text}"

            # Reset deeper levels
            next_level = add_level + 1
            while next_level in self.numbered_headers:
                self.numbered_headers[next_level] = 0
                next_level += 1

            # Scan upper levels
            previous_level = add_level - 1
            while previous_level in self.numbered_headers:
                # MSWord convention: no empty sublevels
                # I.e., sub-sub section (2.0.1) without a sub-section (2.1)
                # is processed as 2.1.1
                if self.numbered_headers[previous_level] == 0:
                    self.numbered_headers[previous_level] += 1

                text = f"{self.numbered_headers[previous_level]}.{text}"
                previous_level -= 1

        self.parents[current_level] = doc.add_heading(
            parent=self.parents[parent_level],
            text=text,
            level=add_level,
        )
        return

    def _add_list_item(
        self,
        *,
        doc: DoclingDocument,
        numid: int,
        ilevel: int,
        elements: list,
        is_numbered: bool = False,
    ) -> None:
        enum_marker = ""

        level = self._get_level()
        prev_indent = self._prev_indent()
        if self._prev_numid() is None:  # Open new list
            self.level_at_new_list = level

            self.parents[level] = doc.add_group(
                label=GroupLabel.LIST, name="list", parent=self.parents[level - 1]
            )

            # Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[level],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )

        elif (
            self._prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and prev_indent < ilevel
        ):  # Open indented list
            for i in range(
                self.level_at_new_list + prev_indent + 1,
                self.level_at_new_list + ilevel + 1,
            ):
                # Determine if this is an unordered list or an ordered list.
                # Set GroupLabel.ORDERED_LIST when it fits.
                self.listIter = 0
                if is_numbered:
                    self.parents[i] = doc.add_group(
                        label=GroupLabel.ORDERED_LIST,
                        name="list",
                        parent=self.parents[i - 1],
                    )
                else:
                    self.parents[i] = doc.add_group(
                        label=GroupLabel.LIST, name="list", parent=self.parents[i - 1]
                    )

            # TODO: Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True

            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[self.level_at_new_list + ilevel],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )
        elif (
            self._prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and ilevel < prev_indent
        ):  # Close list
            for k, v in self.parents.items():
                if k > self.level_at_new_list + ilevel:
                    self.parents[k] = None

            # TODO: Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[self.level_at_new_list + ilevel],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )
            self.listIter = 0

        elif self._prev_numid() == numid or prev_indent == ilevel:
            # TODO: Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[level - 1],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                # Add the list item to the parent group
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )
        return

    def _handle_tables(
        self,
        element: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> None:
        table: Table = Table(element, docx_obj)
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        _log.debug(f"Table grid with {num_rows} rows and {num_cols} columns")

        if num_rows == 1 and num_cols == 1:
            cell_element = table.rows[0].cells[0]
            # In case we have a table of only 1 cell, we consider it furniture
            # And proceed processing the content of the cell as though it's in the document body
            self._walk_linear(cell_element._element, docx_obj, doc)
            return

        data = TableData(num_rows=num_rows, num_cols=num_cols)
        cell_set: set[CT_Tc] = set()
        for row_idx, row in enumerate(table.rows):
            _log.debug(f"Row index {row_idx} with {len(row.cells)} populated cells")
            col_idx = 0
            while col_idx < num_cols:
                cell: _Cell = row.cells[col_idx]
                _log.debug(
                    f" col {col_idx} grid_span {cell.grid_span} grid_cols_before {row.grid_cols_before}"
                )
                if cell is None or cell._tc in cell_set:
                    _log.debug(f"  skipped since repeated content")
                    col_idx += cell.grid_span
                    continue
                else:
                    cell_set.add(cell._tc)

                spanned_idx = row_idx
                spanned_tc: Optional[CT_Tc] = cell._tc
                while spanned_tc == cell._tc:
                    spanned_idx += 1
                    spanned_tc = (
                        table.rows[spanned_idx].cells[col_idx]._tc
                        if spanned_idx < num_rows
                        else None
                    )
                _log.debug(f"  spanned before row {spanned_idx}")

                table_cell = TableCell(
                    text=cell.text,
                    row_span=spanned_idx - row_idx,
                    col_span=cell.grid_span,
                    start_row_offset_idx=row.grid_cols_before + row_idx,
                    end_row_offset_idx=row.grid_cols_before + spanned_idx,
                    start_col_offset_idx=col_idx,
                    end_col_offset_idx=col_idx + cell.grid_span,
                    column_header=row.grid_cols_before + row_idx == 0,
                    row_header=False,
                )
                data.table_cells.append(table_cell)
                col_idx += cell.grid_span

        level = self._get_level()
        doc.add_table(data=data, parent=self.parents[level - 1])
        return

    def _handle_pictures(
        self, docx_obj: DocxDocument, drawing_blip: Any, doc: DoclingDocument
    ) -> None:
        def get_docx_image(drawing_blip):
            rId = drawing_blip[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if rId in docx_obj.part.rels:
                # Access the image part using the relationship ID
                image_part = docx_obj.part.rels[rId].target_part
                image_data = image_part.blob  # Get the binary image data
            return image_data

        level = self._get_level()
        # Open the BytesIO object with PIL to create an Image
        try:
            image_data = get_docx_image(drawing_blip)
            image_bytes = BytesIO(image_data)
            pil_image = Image.open(image_bytes)
            doc.add_picture(
                parent=self.parents[level - 1],
                image=ImageRef.from_pil(image=pil_image, dpi=72),
                caption=None,
            )
        except (UnidentifiedImageError, OSError) as e:
            _log.warning("Warning: image cannot be loaded by Pillow")
            doc.add_picture(
                parent=self.parents[level - 1],
                caption=None,
            )
        return
